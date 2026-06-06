"""Render letter DOCX from either a rich-text body or an uploaded DOCX template.

Placeholders use the {{name}} syntax. Supports:
  - Simple text placeholders (e.g. {{employee_name}})
  - Special block {{salary_breakup_table}} replaced by a formatted table
"""
from __future__ import annotations
from io import BytesIO
from datetime import date
import re
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate
from bs4 import BeautifulSoup

from lib.salary import format_inr


# ------------ Placeholder context ------------

def build_context(company: dict, employee: dict, breakup: dict | None, extras: dict | None = None) -> dict:
    today = date.today()
    ctx = {
        # Company
        "company_name": company.get("name") or "",
        "company_legal_name": company.get("legal_name") or company.get("name") or "",
        "company_address": company.get("address") or "",
        "company_city": company.get("city") or "",
        "company_state": company.get("state") or "",
        "company_pincode": company.get("pincode") or "",
        "company_phone": company.get("phone") or "",
        "company_email": company.get("email") or "",
        "company_website": company.get("website") or "",
        "company_cin": company.get("registration_no") or "",
        "signatory_name": company.get("signatory_name") or "",
        "signatory_designation": company.get("signatory_designation") or "",
        "signatory_department": company.get("signatory_department") or "",
        # Employee
        "employee_name": employee.get("full_name") or "",
        "employee_code": employee.get("employee_code") or "",
        "employee_email": employee.get("email") or "",
        "employee_phone": employee.get("phone") or "",
        "employee_address": employee.get("address") or "",
        "designation": employee.get("designation") or "",
        "department": employee.get("department") or "",
        "reporting_to": employee.get("reporting_to") or "",
        "work_location": employee.get("work_location") or "",
        "date_of_joining": _fmt_date(employee.get("date_of_joining")),
        "date_of_birth": _fmt_date(employee.get("date_of_birth")),
        "probation_months": str(employee.get("probation_months") or ""),
        "employment_type": (employee.get("employment_type") or "").title(),
        # Comp
        "ctc_annual": format_inr(employee.get("ctc_annual") or 0),
        "ctc_annual_words": _amount_in_words(employee.get("ctc_annual") or 0),
        # Dates
        "issue_date": today.strftime("%d %B %Y"),
        "today": today.strftime("%d %B %Y"),
    }
    if breakup:
        ctx["gross_annual"] = format_inr(breakup["totals"]["gross_annual"])
        ctx["gross_monthly"] = format_inr(breakup["totals"]["gross_monthly"])
        ctx["net_annual"] = format_inr(breakup["totals"]["net_annual"])
        ctx["net_monthly"] = format_inr(breakup["totals"]["net_monthly"])
    if extras:
        ctx.update({k: ("" if v is None else str(v)) for k, v in extras.items()})
    return ctx


def _fmt_date(d) -> str:
    if not d: return ""
    if isinstance(d, str):
        try:
            from dateutil.parser import parse
            d = parse(d).date()
        except Exception:
            return d
    try:
        return d.strftime("%d %B %Y")
    except Exception:
        return str(d)


def _amount_in_words(n) -> str:
    try:
        n = int(round(float(n)))
    except Exception:
        return ""
    if n == 0: return "Zero Rupees Only"
    # Simple Indian numbering
    units = ["", "One","Two","Three","Four","Five","Six","Seven","Eight","Nine",
             "Ten","Eleven","Twelve","Thirteen","Fourteen","Fifteen","Sixteen",
             "Seventeen","Eighteen","Nineteen"]
    tens = ["", "", "Twenty","Thirty","Forty","Fifty","Sixty","Seventy","Eighty","Ninety"]
    def two(n):
        if n < 20: return units[n]
        return tens[n//10] + ((" " + units[n%10]) if n%10 else "")
    def three(n):
        s = ""
        if n >= 100:
            s += units[n//100] + " Hundred"
            n %= 100
            if n: s += " "
        if n: s += two(n)
        return s
    parts = []
    crore = n // 10000000; n %= 10000000
    lakh = n // 100000; n %= 100000
    thou = n // 1000; n %= 1000
    if crore: parts.append(three(crore) + " Crore")
    if lakh: parts.append(two(lakh) + " Lakh")
    if thou: parts.append(two(thou) + " Thousand")
    if n: parts.append(three(n))
    return " ".join(parts) + " Rupees Only"


# ------------ Render from rich text (HTML) ------------

def render_richtext_to_docx(html_body: str, ctx: dict, company_logo_url: str | None,
                            breakup: dict | None, letter_title: str | None = None) -> bytes:
    """Convert a rich-text template into a DOCX. Tight spacing + logo header on every page."""
    # Substitute placeholders first (text-level)
    body = _substitute(html_body or "", ctx)

    doc = Document()
    # Tight margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Comfortable Normal style — single line spacing, modest space after.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15
    # Also tighten the List Bullet and List Number built-in styles
    for sn in ("List Bullet", "List Number", "List Paragraph"):
        try:
            s = doc.styles[sn]
            s.paragraph_format.space_before = Pt(0)
            s.paragraph_format.space_after = Pt(0)
            s.paragraph_format.line_spacing = 1.0
        except KeyError:
            pass

    # Logo intentionally NOT rendered as page header — keeps the top tight.
    # The {{signature_block}} token below provides the two-column signature.

    # Stash context on the doc so the {{signature_block}} renderer can use it.
    doc._signature_ctx = ctx

    if letter_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(letter_title)
        r.bold = True
        r.font.size = Pt(12)

    # Walk HTML and add to docx
    soup = BeautifulSoup(body, "lxml")
    root = soup.body or soup
    for el in root.children:
        _render_html_node(doc, el, breakup)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _is_empty_html_node(node) -> bool:
    """True if a <p>/<div> contains only whitespace or a lone <br> (Quill empty line)."""
    nm = getattr(node, "name", None)
    if nm not in ("p", "div"):
        return False
    txt = (node.get_text() or "").strip()
    if txt:
        return False
    # Only whitespace -- check if it just has <br> tags
    children = [c for c in node.children if getattr(c, "name", None) or str(c).strip()]
    if not children:
        return True
    if all(getattr(c, "name", None) == "br" for c in children):
        return True
    return False


_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")

def _substitute(text: str, ctx: dict) -> str:
    def repl(m):
        key = m.group(1)
        return str(ctx.get(key, m.group(0)))
    return _PLACEHOLDER_RE.sub(repl, text)


def _render_html_node(doc: Document, node, breakup):
    name = getattr(node, "name", None)
    if name is None:
        # NavigableString
        text = str(node).strip()
        if text:
            doc.add_paragraph(text)
        return

    # Skip Quill empty-line paragraphs (<p><br></p>) — the paragraph
    # space_after on adjacent paragraphs already provides visual separation.
    if _is_empty_html_node(node):
        return

    text_content = node.get_text() if hasattr(node, "get_text") else ""

    # Salary breakup token
    if "{{salary_breakup_table}}" in text_content and breakup:
        _add_breakup_table(doc, breakup)
        return

    # Two-column signature block token
    if "{{signature_block}}" in text_content:
        _add_signature_block(doc)
        return

    if name in ("p", "div"):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.15
        _add_inline(para, node)
    elif name in ("h1","h2","h3","h4","h5","h6"):
        level = int(name[1])
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(node.get_text())
        run.bold = True
        run.font.size = Pt({1:14,2:13,3:12,4:11,5:11,6:11}[level])
    elif name in ("ul","ol"):
        for li in node.find_all("li", recursive=False):
            para = doc.add_paragraph(style="List Bullet" if name == "ul" else "List Number")
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            _add_inline(para, li)
    elif name == "br":
        # standalone <br> outside a paragraph — ignore (we don't want gap-only paragraphs)
        return
    elif name == "table":
        _add_html_table(doc, node)
    else:
        # fallback: treat as paragraph
        if text_content.strip():
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(4)
            _add_inline(para, node)


def _add_inline(para, node):
    for child in node.children:
        nm = getattr(child, "name", None)
        if nm is None:
            run = para.add_run(str(child))
        elif nm in ("strong","b"):
            run = para.add_run(child.get_text())
            run.bold = True
        elif nm in ("em","i"):
            run = para.add_run(child.get_text())
            run.italic = True
        elif nm == "u":
            run = para.add_run(child.get_text())
            run.underline = True
        elif nm == "br":
            para.add_run().add_break()
        else:
            _add_inline(para, child)


def _add_html_table(doc, table_node):
    rows = table_node.find_all("tr")
    if not rows: return
    cols = max(len(r.find_all(["td","th"])) for r in rows)
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = "Light Grid Accent 1"
    for ri, r in enumerate(rows):
        cells = r.find_all(["td","th"])
        for ci, c in enumerate(cells):
            t.rows[ri].cells[ci].text = c.get_text(strip=True)


def _add_breakup_table(doc, breakup: dict):
    rows = breakup["annual"]
    t = doc.add_table(rows=1 + len(rows) + 3, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Monthly (INR)"
    hdr[2].text = "Annual (INR)"
    for r in t.rows[0].cells:
        for p in r.paragraphs:
            for run in p.runs:
                run.bold = True
    for i, (a, m) in enumerate(zip(breakup["annual"], breakup["monthly"]), start=1):
        c = t.rows[i].cells
        c[0].text = a["name"] + ("" if a["type"] == "earning" else "  (Deduction)")
        c[1].text = format_inr(m["amount"])
        c[2].text = format_inr(a["amount"])
    totals = breakup["totals"]
    idx = len(rows) + 1
    t.rows[idx].cells[0].text = "Gross"
    t.rows[idx].cells[1].text = format_inr(totals["gross_monthly"])
    t.rows[idx].cells[2].text = format_inr(totals["gross_annual"])
    t.rows[idx+1].cells[0].text = "Deductions"
    t.rows[idx+1].cells[2].text = format_inr(totals["deductions_annual"])
    t.rows[idx+2].cells[0].text = "Net Take-home"
    t.rows[idx+2].cells[1].text = format_inr(totals["net_monthly"])
    t.rows[idx+2].cells[2].text = format_inr(totals["net_annual"])
    for ri in (idx, idx+1, idx+2):
        for cell in t.rows[ri].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True


def _add_signature_block(doc: Document):
    """
    Render a two-column, borderless signature block.

    Left column:                     Right column:
    For {{company_name}},            Received and Accepted
    (blank)
    {{signatory_name}}               Name & Signature: ______________
    {{signatory_designation}}        Date: ______________
    """
    # We render the substituted values via a separate path — the renderer
    # already substituted {{company_name}} etc. in the surrounding text,
    # but this token is replaced as a block, so we need to read context.
    # Use the doc.signature_ctx attribute populated by the caller.
    ctx = getattr(doc, "_signature_ctx", {})
    company_name = ctx.get("company_name", "")
    sig_name = ctx.get("signatory_name", "")
    sig_des = ctx.get("signatory_designation", "")

    # Small spacer above
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(0)
    sp.add_run("").font.size = Pt(8)

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    # Remove all borders
    from docx.oxml.ns import qn
    tbl_pr = t._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)

    left, right = t.rows[0].cells
    # Make columns share equal width
    for cell in (left, right):
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15

    def _add_line(cell, text, bold=False, top=False):
        p = cell.paragraphs[0] if top and not cell.paragraphs[0].runs else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        if bold:
            run.bold = True

    # Left column
    left_para = left.paragraphs[0]
    left_para.paragraph_format.space_after = Pt(2)
    r = left_para.add_run(f"For {company_name},")
    r.bold = True
    # blank line for signature space
    sp_l = left.add_paragraph(""); sp_l.add_run("").font.size = Pt(20)
    _add_line(left, sig_name, bold=True)
    _add_line(left, sig_des)

    # Right column
    right_para = right.paragraphs[0]
    right_para.paragraph_format.space_after = Pt(2)
    r = right_para.add_run("Received and Accepted")
    r.bold = True
    # blank for spacing
    sp_r = right.add_paragraph(""); sp_r.add_run("").font.size = Pt(20)
    _add_line(right, "Name & Signature: ______________________")
    _add_line(right, "")  # tiny gap
    _add_line(right, "Date: ______________________")


# ------------ Render from uploaded DOCX template ------------

def render_docx_template(template_bytes: bytes, ctx: dict, breakup: dict | None) -> bytes:
    """Use docxtpl ({{ jinja }} placeholders) for uploaded DOCX templates."""
    tpl = DocxTemplate(BytesIO(template_bytes))
    jinja_ctx = dict(ctx)
    if breakup:
        jinja_ctx["salary_annual"] = breakup["annual"]
        jinja_ctx["salary_monthly"] = breakup["monthly"]
        jinja_ctx["salary_totals"] = breakup["totals"]
    tpl.render(jinja_ctx)
    out = BytesIO()
    tpl.save(out)
    return out.getvalue()
